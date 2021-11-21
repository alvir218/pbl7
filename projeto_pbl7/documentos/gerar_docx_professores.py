import requests
from docx import Document
from docx.shared import Inches


#exemplo da estrutura json da aplicacao professores
# {
#     "cod_professor": "cfd38e58-82c9-4d09-8e9c-054a6d07efe6",
#     "nome_professor": "nomeprofessor2",
#     "ru": "ru2",
#     "matricula": 2,
#     "formacao": "formacao2",
#     "titulacao": "titulacao2",
#     "area_titulacao": "areatitulacao2",
#     "endereco_completo": "enderecocompleto2",
#     "telefone_fixo": "2222222",
#     "whatsapp": "2222222",
#     "email": "email2@email2.com",
#     "data_ingresso_instituicao": "2021-11-20",
#     "data_docencia": null,
#     "data_docencia_instituicao": null,
#     "funcao": "funcao2",
#     "regime_trabalho": "regimetrabalho2"
# }


documento = Document()
documento.add_heading('Professores', 0)

paragrafo = documento.add_paragraph('Lista de Professores')
paragrafo.add_run(' - Documento gerado por meio de Software API').italic = True







print('\n\n   ensaio3 convertendo request pra json\n\n')

#ids_professores = Professores.objects.values_list('cod_professor', flat=True)
#ids_professores = professores.values_list('cod_professor', flat=True)
professores_obtidos = requests.get('http://127.0.0.1:8000/professores/')

print('\n\nlistando ids dos professores 3\n')
#print(f'id: {professores_obtidos.json().get("cod_professor")}')
#print(professores_obtidos.get('total'))
#print(professores_obtidos['cod_professor'])
json_professores = professores_obtidos.json()
print(json_professores)


print('\n\n   ensaio 4 trabalhando todos os registros\n\n')

for elemento in json_professores:

    print('\n\n')

    cod_professor = elemento.get('cod_professor')

    request = requests.get('http://127.0.0.1:8000/professores/' + cod_professor)

    dados_da_request = request.json()

    print('nome do professor: {}'.format(dados_da_request['nome_professor']))
    documento.add_heading(dados_da_request['nome_professor'], level=1)

    
    print('ru: '+ dados_da_request['ru'])
    documento.add_paragraph('RU-> '+dados_da_request['ru'], style='Intense Quote')


    print('matricula:', dados_da_request['matricula'])
    documento.add_paragraph('Matricula-> '+str(dados_da_request['matricula']), style='Intense Quote')


    print('formacao: {}'.format(dados_da_request['formacao']))
    documento.add_paragraph('Formacao-> '+dados_da_request['formacao'], style='Intense Quote')


    print('titulacao: {}'.format(dados_da_request['titulacao']))
    documento.add_paragraph('Titulacao-> '+dados_da_request['titulacao'], style='Intense Quote')


    print('area da titulacao: {}'.format(dados_da_request['area_titulacao']))
    documento.add_paragraph('Area da Titulacao-> '+dados_da_request['area_titulacao'], style='Intense Quote')


    print('endereco completo: {}'.format(dados_da_request['endereco_completo']))
    documento.add_paragraph('Endereco Completo-> '+dados_da_request['endereco_completo'], style='Intense Quote')
    

    print('telefone fixo: {}'.format(dados_da_request['telefone_fixo']))
    documento.add_paragraph('Telefone Fixo-> '+dados_da_request['telefone_fixo'], style='Intense Quote')
    

    print('whatsapp: {}'.format(dados_da_request['whatsapp']))
    documento.add_paragraph('WhatsApp-> '+dados_da_request['whatsapp'], style='Intense Quote')
    

    print('email: {}'.format(dados_da_request['email']))
    documento.add_paragraph('E-mail-> '+dados_da_request['email'], style='Intense Quote')
    

    print('data de ingresso na instituicao: {}'.format(dados_da_request['data_ingresso_instituicao']))
    documento.add_paragraph('Data Ingresso-> '+dados_da_request['data_ingresso_instituicao'], style='Intense Quote')
    

    if not dados_da_request['data_docencia'] == None:
        print('docente desde: {}'.format(dados_da_request['data_docencia']))
        documento.add_paragraph('Docente desde '+dados_da_request['data_docencia'], style='Intense Quote').italic = True
        

    if not dados_da_request['data_docencia_instituicao'] == None:
        print('docente na instituicao desde: {}'.format(dados_da_request['data_docencia_instituicao']))
        documento.add_paragraph('Docente na Instituicao desde '+dados_da_request['data_docencia_instituicao'], style='Intense Quote').italic = True
        

    print('funcao: {}'.format(dados_da_request['funcao']))
    documento.add_paragraph('Funcao-> '+dados_da_request['funcao'], style='Intense Quote').bold = True
    

    print('regime de trabalho: {}'.format(dados_da_request['regime_trabalho']))
    documento.add_paragraph('Regime de trabalho-> '+dados_da_request['regime_trabalho'], style='Intense Quote').bold = True

documento.save('Professores.docx')
