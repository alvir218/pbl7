from docx import Document
from docx.shared import Inches

professores = ("Amilton", "Ederson", "Fernanda", "Vinicius", "Felipe")
materias = ("Redes", "Logica Programavel", "Fisica", "Programação", "Eletricidade")
count = 0

document = Document()
document.add_heading('Professores', 0)

p = document.add_paragraph('Lista de ')
p.add_run('Professores').bold = True
p.add_run(' (Teste)').italic = True

for i in professores:
        document.add_heading(i, level=1)
        document.add_paragraph(materias[count], style='Intense Quote')
        count+=1
document.add_page_break()

document.save('teste.docx')
